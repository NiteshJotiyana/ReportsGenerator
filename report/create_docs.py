from docx import  Document
from docx.shared import Inches
from PIL import Image


def create_empty_doc(file_name):
    doc = Document()
    doc.save(file_name)
    print('Word document created successfully.')


def create_doc_with_data(file_name:str,doc_heading:str,image_files):
    doc = Document()
    # Add a title to the document
    doc.add_heading(doc_heading, 0)

    # Example image filenames (replace with your own image paths)
    #image_files = ['screenshots/tab_0_chrome_screenshot_dummy.png']

    # Iterate over each image file and add it to the document
    for image_file in image_files:
        # Add a heading for each image
        doc.add_heading(f'Image: {image_file}', level=1)

        # Add the image to the document
        doc.add_picture(image_file, width=Inches(8))  # Adjust width as needed

        # Add a paragraph after each image (optional)
        doc.add_paragraph('')  # Add an empty paragraph

        # for paragraph in doc.paragraphs:
        #     if paragraph.style.name.startswith('Heading'):
        #         print("heading: "+paragraph.text)

    # Save the document
    doc.save(file_name)

    print('Word document created successfully.')
